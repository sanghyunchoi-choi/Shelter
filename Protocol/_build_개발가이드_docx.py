# -*- coding: utf-8 -*-
"""Smart Shelter 개발가이드 Word 문서 생성.

수정: Protocol/개발가이드.md 편집 후
실행: python _build_개발가이드_docx.py
출력: Protocol/개발가이드.docx
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).parent
MD_SRC = HERE / "개발가이드.md"
OUT = HERE / "개발가이드.docx"
OUT_V1 = HERE / "개발가이드_v1.0.docx"
REV = "1.3"
BUILT = date.today().isoformat()

FONT_KO = "맑은 고딕"
FONT_MONO = "Consolas"
COLOR_TITLE = RGBColor(0x1F, 0x4E, 0x79)


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = FONT_KO
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = FONT_KO
        hs.font.color.rgb = COLOR_TITLE
        hs._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
        hs.font.bold = True
        if level == 1:
            hs.font.size = Pt(16)
        elif level == 2:
            hs.font.size = Pt(13)
        else:
            hs.font.size = Pt(11)


def add_para(doc, text, bold=False, mono=False, size=10, space_after=6, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = FONT_MONO if mono else FONT_KO
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_MONO if mono else FONT_KO)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = FONT_KO
                r.font.size = Pt(9)
                r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = FONT_KO
                    r.font.size = Pt(9)
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text.strip())
    run.font.name = FONT_MONO
    run.font.size = Pt(8.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_MONO)


def add_diagram(doc, text):
    add_code_block(doc, text)


def build_cover(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Smart Shelter\n개발 가이드")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = COLOR_TITLE
    r.font.name = FONT_KO
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        f"제어보드 · 전원보드 — 서버 연동 개발자용\n\n"
        f"문서 버전 rev {REV}  |  생성일 {BUILT}\n"
        f"대상: MOS_version (MQTT 1883)"
    )
    sr.font.size = Pt(11)
    sr.font.name = FONT_KO
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    doc.add_page_break()


def build_body(doc):
    doc.add_heading("1. 이 문서의 목적", level=1)
    add_para(
        doc,
        "본 문서는 Smart Shelter 제어보드(Control Board)와 전원보드(Power Board, 8ch)를 "
        "납품받아 자체 관제 서버·애플리케이션을 개발하는 엔지니어를 위한 진입점(Entry Guide)입니다.",
    )
    add_para(
        doc,
        "전기제품 구매 시 동봉되는 사용설명서와 같이, 본 가이드와 관련 규격 문서를 참조하면 "
        "MQTT 기반 관제 서버를 독립적으로 구현할 수 있습니다. §8에 tele/cmnd/stat 복사용 예시 포함.",
    )
    add_table(
        doc,
        ["구분", "내용"],
        [
            ("본 문서", "시스템 구성, Quick Start, uid/b_id, 체크리스트"),
            ("Shelter_Protocol.xlsx", "MQTT 토픽·JSON 페이로드 정본"),
            ("Shelter_Installation_Guide.md", "현장·공장 설치·IP 설정"),
            ("Shelter_CMS_Deploy", "CMS 참조 배포 (선택)"),
        ],
        col_widths=[4, 12],
    )

    doc.add_heading("2. 문서 체계 (납품 패키지)", level=1)
    add_table(
        doc,
        ["문서", "파일", "대상"],
        [
            ("개발 가이드 (본 문서)", "개발가이드.docx / 개발가이드_v1.0.docx", "서버·연동 개발자"),
            ("MQTT 프로토콜 규격", "Shelter_Protocol.xlsx / _v1.0", "서버 개발 정본"),
            ("MQTT 프로토콜 상세", "Shelter_Protocol.md", "개발자·변경 이력"),
            ("설치 가이드", "Shelter_Installation_Guide.md", "설치 업체"),
            ("현장 운영·배포", "Shelter_Ops_Deploy.xlsx", "관제 PC·공무원"),
        ],
        col_widths=[4.5, 5, 5],
    )
    add_para(doc, "개발 시 우선순위: Quick Start(§6) → Shelter_Protocol.xlsx → Shelter_Protocol.md 상세.", bold=True)

    doc.add_heading("3. 시스템 아키텍처", level=1)
    doc.add_heading("3.1 구성 요소", level=2)
    add_diagram(
        doc,
        """
┌─────────────┐     MQTT 1883      ┌──────────────────┐
│  제어보드    │◄──────────────────►│  귀사 관제 서버   │
│ (STM32+ETH) │   dev/tele/*       │  (Mosquitto 등)  │
│             │   dev/stat/*       │                  │
│             │   dev/cmnd/*       └──────────────────┘
└──────┬──────┘
       │ RS485
       ▼
┌─────────────┐
│  전원보드    │  8ch 전력 ON/OFF · 전류 측정 · DIP ID
└─────────────┘
""",
    )
    doc.add_heading("3.2 역할 요약", level=2)
    add_table(
        doc,
        ["구성요소", "역할", "서버가 직접 다루는가?"],
        [
            ("제어보드", "센서·릴레이·PB 집계, MQTT 게이트웨이", "예 — MQTT만 구현"),
            ("전원보드", "8ch 전력, RS485 슬레이브", "아니오 — tele/pb, cmnd/pb"),
            ("관제 서버", "브로커 + 구독/발행 + UI", "귀사 개발"),
        ],
        col_widths=[3, 7, 5],
    )
    doc.add_heading("3.3 연동 경계", level=2)
    add_para(doc, "서버 엔지니어는 제어보드 MQTT API만 구현하면 됩니다.", bold=True)
    add_para(doc, "• 구독: dev/tele/#, dev/stat/#")
    add_para(doc, "• 발행: dev/cmnd/{device}")
    add_para(doc, "• RS485·Modbus·GPIO는 제어보드 펌웨어가 처리합니다.")

    doc.add_heading("4. 하드웨어 개요", level=1)
    doc.add_heading("4.1 제어보드", level=2)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ("통신", "이더넷 → MQTT 포트 1883"),
            ("장치 식별", "dev/tele/state 의 uid (STM32 96-bit, 24자리 HEX)"),
            ("최초 설정", "PC 직결 http://192.168.0.100"),
            ("현장 LAN", "Flash DHCP(권장) 또는 STATIC"),
        ],
        col_widths=[4, 12],
    )
    doc.add_heading("4.2 전원보드 (8ch)", level=2)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ("기능", "8ch ON/OFF, 채널별 전류·전력"),
            ("연결", "RS485 (제어보드 ↔ PB)"),
            ("보드 ID", "DIP 4bit → b_id 0~15 (고정값 아님)"),
            ("MQTT", "dev/tele/pb, dev/cmnd/pb"),
        ],
        col_widths=[4, 12],
    )

    doc.add_heading("5. 장치 식별", level=1)
    doc.add_heading("5.1 제어보드 — uid", level=2)
    add_para(doc, "토픽 dev/tele/state — ONLINE 시 uid 필드로 쉘터 1대를 식별합니다.")

    doc.add_heading("5.2 전원보드 — b_id", level=2)
    add_para(doc, "b_id는 고정값이 아닙니다. 전원보드 DIP 4bit(0~15) 실측값입니다.", bold=True)
    add_table(
        doc,
        ["단계", "동작"],
        [
            ("전원보드", "DIP → RS485 패킷 byte[1]"),
            ("제어보드", "RS485 수신 → b_id 학습"),
            ("MQTT tele/stat", "학습한 b_id 전달"),
            ("서버 cmnd", "tele/get_pb와 동일 b_id 필수"),
        ],
        col_widths=[3.5, 12],
    )
    add_table(
        doc,
        ["스위치", "bit", "ON(LOW) 시"],
        [
            ("ID1", "bit0 (LSB)", "+1"),
            ("ID2", "bit1", "+2"),
            ("ID3", "bit2", "+4"),
            ("ID4", "bit3 (MSB)", "+8"),
        ],
        col_widths=[3, 4, 8],
    )
    add_para(doc, "예: ID1만 ON → b_id=1 / ID1+ID3 → b_id=5 / 전부 OFF → b_id=0")
    add_para(doc, "서버: b_id 하드코딩 금지. tele/pb 또는 get_pb 응답을 캐시해 사용 (CMS 초기값 0).", bold=True)

    doc.add_heading("5.3 cmnd 복합 JSON (v2.4)", level=2)
    add_para(doc, "AC·AP는 한 JSON에 pwr+mode+temp 등 여러 필드를 동시에 넣을 수 있습니다.")
    add_code_block(doc, '{"pwr":1,"mode":"COOL","temp":24,"spd":2}\n{"set_ac":{"pwr":1,"mode":"COOL","temp":24,"spd":2}}')

    doc.add_heading("6. Quick Start — MQTT 연동", level=1)
    add_para(doc, "사전: 제어보드 LAN·브로커 IP 설정 완료, Mosquitto 1883 수신 대기.")

    doc.add_heading("Step 1 — ONLINE 확인", level=2)
    add_code_block(doc, 'mosquitto_sub -h <브로커IP> -p 1883 -t "dev/tele/state" -v')

    doc.add_heading("Step 2 — PB b_id 확인", level=2)
    add_code_block(doc, 'mosquitto_sub -h <브로커IP> -p 1883 -t "dev/tele/pb" -v')
    add_code_block(
        doc,
        '{"b_id":3,"pb":[{"ch":1,"sw":1,"c":1.14,"w":250.8,"v":220}, ...],"conn":1}',
    )
    add_para(doc, "b_id 값을 서버에 저장합니다.")

    doc.add_heading("Step 3 — PB 채널 제어", level=2)
    add_code_block(
        doc,
        'mosquitto_pub -h <브로커IP> -t "dev/cmnd/pb" -m \'{"b_id":3,"ch":1,"sw":1}\'',
    )
    add_para(doc, "dev/stat/pb 에서 result:1 확인.")

    doc.add_heading("Step 4 — 전체 동기화", level=2)
    add_code_block(
        doc,
        'mosquitto_pub -h <브로커IP> -t "dev/cmnd/dev" -m \'{"data":"ALL_DATA"}\'',
    )

    doc.add_heading("7. MQTT 토픽 요약", level=1)
    add_table(
        doc,
        ["Prefix", "방향", "설명"],
        [
            ("dev/tele/…", "장치→서버", "Telemetry"),
            ("dev/stat/…", "장치→서버", "명령 결과 (result)"),
            ("dev/cmnd/…", "서버→장치", "제어 명령"),
        ],
        col_widths=[4, 3.5, 8],
    )
    add_table(
        doc,
        ["토픽", "장치", "주기"],
        [
            ("dev/tele/state", "시스템", "연결/해제"),
            ("dev/tele/pb", "전원보드 8ch", "10초"),
            ("dev/tele/ad", "릴레이 15ch", "10분"),
            ("dev/tele/ac", "에어컨 LG", "10분"),
            ("dev/tele/ap", "공기청정기", "10분"),
            ("dev/tele/fan", "내부 팬", "10분"),
            ("dev/tele/dust", "미세먼지", "10분"),
            ("dev/tele/th_in", "실내 온습도", "10분"),
            ("dev/tele/th_out", "실외 온습도", "10분"),
            ("dev/tele/input", "외부 입력 4ch", "10분"),
            ("dev/cmnd/dev", "시스템", "ALL_DATA 등"),
        ],
        col_widths=[5, 4, 3],
    )
    add_para(doc, "전체 payload·명령: Shelter_Protocol.xlsx (정본).")

    doc.add_heading("8. 프로토콜 예시 및 테스트", level=1)
    add_para(
        doc,
        "아래 JSON은 복사해서 바로 테스트할 수 있는 예시입니다. "
        "<브로커IP>는 Mosquitto IP, b_id 예시는 3 (tele/pb 실측값으로 교체).",
    )
    doc.add_heading("8.1 전체 구독", level=2)
    add_code_block(
        doc,
        'mosquitto_sub -h <브로커IP> -p 1883 -t "dev/tele/#" -t "dev/stat/#" -v',
    )

    doc.add_heading("8.2 state", level=2)
    add_table(
        doc,
        ["구분", "토픽", "Payload 예시"],
        [
            ("tele ONLINE", "dev/tele/state", '{"status":"ONLINE","uid":"0008DC77631F","ip":"192.168.0.115","mode":"DHCP","conn":1}'),
            ("tele OFFLINE", "dev/tele/state", '{"status":"OFFLINE"}'),
        ],
        col_widths=[2.5, 4, 9],
    )

    doc.add_heading("8.3 전원보드 pb", level=2)
    add_para(doc, "tele dev/tele/pb (연결):", bold=True)
    add_code_block(
        doc,
        '{"b_id":3,"pb":[{"ch":1,"sw":1,"c":1.14,"w":250.8,"v":220},{"ch":2,"sw":1,"c":1.04,"w":228.5,"v":220},'
        '... ch3~8 ...],"conn":1}',
    )
    add_para(doc, "tele 미연결: {\"pb\":0,\"conn\":0}")
    add_table(
        doc,
        ["cmnd dev/cmnd/pb", "Payload"],
        [
            ("채널1 ON", '{"b_id":3,"ch":1,"sw":1}'),
            ("채널1 OFF", '{"b_id":3,"ch":1,"sw":0}'),
            ("8ch 일괄", '{"b_id":3,"set_pb":"11001101"}'),
            ("조회", '{"b_id":3,"get_pb":"state"}'),
        ],
        col_widths=[3, 12],
    )
    add_para(doc, "stat 성공: tele 동일 + \"result\":1")
    add_para(doc, "stat 실패: {\"b_id\":3,\"pb\":0,\"conn\":1,\"result\":0}")
    add_code_block(
        doc,
        'mosquitto_pub -h <브로커IP> -t "dev/cmnd/pb" -m \'{"b_id":3,"get_pb":"state"}\'\n'
        'mosquitto_pub -h <브로커IP> -t "dev/cmnd/pb" -m \'{"b_id":3,"ch":1,"sw":1}\'',
    )

    doc.add_heading("8.4 AD 릴레이 15ch", level=2)
    add_para(doc, "tele: {\"relays\":[1,0,0,...15개],\"conn\":1}")
    add_table(
        doc,
        ["cmnd", "Payload"],
        [
            ("CH1 ON", '{"ch":1,"relay_1":1}'),
            ("CH1 OFF", '{"ch":1,"relay_1":0}'),
            ("15ch 일괄", '{"set_ad":"100000000000000"}'),
            ("조회", '{"get_ad":"state"}'),
        ],
        col_widths=[3, 12],
    )
    add_para(doc, "stat: {\"relays\":[...],\"conn\":1,\"result\":1}")

    doc.add_heading("8.5 fan / 8.6 ac / 8.7 ap", level=2)
    add_table(
        doc,
        ["장치", "tele 예", "cmnd 예"],
        [
            ("fan", '{"mode":"MANUAL","duty":50,"rpm":1234,"conn":1}', '{"set_fan":{"mode":"MANUAL","duty":50}}'),
            ("ac", '{"pwr":1,"mode":"COOL","temp":24,"curr":26.0,"spd":3,"err":0,"conn":1}', '{"set_ac":{...}} 또는 {"pwr":1,"mode":"COOL","temp":24,"spd":2}'),
            ("ap", '{"pwr":1,"mode":"AUTO","spd":3,"temp_out":18.2,"tvoc":120,...,"conn":1}', '{"set_ap":{...}} 또는 {"pwr":1,"mode":"AUTO","spd":2}'),
        ],
        col_widths=[2, 6.5, 6.5],
    )

    doc.add_heading("8.8 환경 센서 (tele 전용)", level=2)
    add_table(
        doc,
        ["토픽", "연결", "미연결"],
        [
            ("dev/tele/dust", '{"pm1_0":15,"pm2_5":24,"pm10":24,"conn":1}', '{"dust":0,"conn":0}'),
            ("dev/tele/th_in", '{"temp":25.3,"humi":42.1,"conn":1}', '{"thindoor":0,"conn":0}'),
            ("dev/tele/th_out", '{"temp":12.5,"humi":35.0,"conn":1}', '{"thoutdoor":0,"conn":0}'),
        ],
        col_widths=[4, 5.5, 5.5],
    )

    doc.add_heading("8.9 input / 8.10 dev", level=2)
    add_para(doc, "input tele: {\"in\":[1,0,0,1],\"count\":[12,0,0,3],\"conn\":1}")
    add_para(doc, "input cmnd: {\"get_input\":\"state\"} / {\"reset_count\":1}")
    add_table(
        doc,
        ["dev/cmnd/dev", "Payload", "stat 예"],
        [
            ("ALL_DATA", '{"data":"ALL_DATA"}', '{"device":"ALL_DATA","conn":1,"result":1}'),
            ("재부팅", '{"reset":"DEV_RESET"}', '{"device":"DEV_RESET","conn":0,"result":1}'),
            ("브로커", '{"set_broker":{"ip":[192,168,0,107],"port":1883}}', '{"device":"SET_BROKER","conn":1,"result":1}'),
        ],
        col_widths=[3, 7, 5],
    )

    doc.add_heading("8.11 테스트 시나리오", level=2)
    add_table(
        doc,
        ["#", "동작", "확인"],
        [
            ("1", "mosquitto_sub tele/# stat/#", "구독"),
            ("2", "제어보드 전원 ON", "state ONLINE + uid"),
            ("3", "cmnd/dev ALL_DATA", "각 tele 1회"),
            ("4", "tele/pb b_id 기록", "예: 3"),
            ("5", "cmnd/pb get_pb", "stat result:1"),
            ("6", "cmnd/pb ch:1 sw:1", "tele ch1 ON"),
            ("7", "cmnd/ad get_ad", "relays 15ch"),
        ],
        col_widths=[1, 6, 8],
    )

    doc.add_page_break()
    doc.add_heading("9. 연동 체크리스트", level=1)
    doc.add_heading("9.1 서버 개발 완료 기준", level=2)
    checks = [
        "MQTT 브로커 1883 수신·발행",
        "dev/tele/state ONLINE/OFFLINE (uid 등록)",
        "dev/tele/pb 구독 → b_id 캐시 (초기 0, 하드코딩 금지)",
        "tele/pb — pb[] 배열로 파싱 (result 불필요)",
        "dev/cmnd/pb — 캐시 b_id로 ch / set_pb / get_pb",
        "dev/stat/pb — result, conn 처리",
        "conn:0 시 UI 오프라인 (메인 ONLINE ≠ 하위 전부 연결)",
        'dev/cmnd/dev {"data":"ALL_DATA"} 초기 동기화',
        "(선택) ad, ac, ap, fan, dust, th_in, th_out, input",
    ]
    for c in checks:
        add_para(doc, f"☐  {c}")

    doc.add_heading("9.2 현장 연동 확인", level=2)
    for c in [
        "dev/tele/state ONLINE",
        "dev/tele/pb conn:1, pb[] 8채널",
        "PB 제어 후 stat result:1",
        "b_id 불일치 cmnd → result:0 (정상 거부)",
    ]:
        add_para(doc, f"☐  {c}")

    doc.add_heading("10. 동작 규칙 (요약)", level=1)
    add_table(
        doc,
        ["규칙", "설명"],
        [
            ("10분 tele", "장치별 10분 (PB만 10초)"),
            ("명령 후 stat", "cmnd → stat 발행 + 타이머 리셋"),
            ("재연결", "복구 후 tele 1회 + ALL_DATA"),
            ("LWT", "비정상 종료 → state OFFLINE"),
            ("PB conn", "RS485 10회 실패 → conn:0"),
            ("복합 cmnd", "AC/AP 한 JSON 다필드 (v2.4)"),
        ],
        col_widths=[4, 12],
    )

    doc.add_heading("11. 납품 범위 (일반)", level=1)
    add_table(
        doc,
        ["포함 (예)", "미포함 (계약 별도)"],
        [
            ("펌웨어 HEX/ELF", "귀사 관제 서버·UI"),
            ("MQTT 규격 xlsx/md", "CMS 소스"),
            ("설치·개발 가이드", "펌웨어 소스 (별도 계약)"),
            ("CMS 배포 패키지 (선택)", "Himpel RS485 / LG Modbus 상세"),
        ],
        col_widths=[7, 8],
    )

    doc.add_page_break()
    doc.add_heading("부록 A — 전원보드 RS485 (참고)", level=1)
    add_para(
        doc,
        "대부분의 서버 개발자에게 불필요. 제어보드 없이 PB만 직접 RS485 연동할 때 참고.",
        bold=True,
    )
    doc.add_heading("A.1 보고 패킷 (PB→제어보드, 30 byte)", level=2)
    add_table(
        doc,
        ["Offset", "내용"],
        [
            ("0", "STX 0x02"),
            ("1", "Board ID (DIP 0~15)"),
            ("2", "Length 0x1E"),
            ("3", "Command 0xBB"),
            ("4~27", "8ch × 3byte (상태 + 전류 mA)"),
            ("28", "ETX 0x03"),
            ("29", "BCC (0~28 XOR)"),
        ],
        col_widths=[3, 12],
    )
    doc.add_heading("A.2 제어 패킷 (제어보드→PB, 14 byte)", level=2)
    add_table(
        doc,
        ["Offset", "내용"],
        [
            ("0", "STX 0x02"),
            ("1", "Board ID"),
            ("2", "Length 0x0E"),
            ("3", "Command 0xAA"),
            ("4~11", "8ch 상태"),
            ("12", "ETX 0x03"),
            ("13", "BCC"),
        ],
        col_widths=[3, 12],
    )
    add_para(doc, "PB는 자신의 ID와 일치하는 제어 패킷만 수신합니다.")

    doc.add_heading("부록 B — 관련 문서", level=1)
    add_table(
        doc,
        ["문서", "경로"],
        [
            ("MQTT 규격 (엑셀)", "Protocol/Shelter_Protocol.xlsx · _v1.0.xlsx"),
            ("MQTT 규격 (MD)", "Protocol/Shelter_Protocol.md"),
            ("설치 가이드", "Protocol/Shelter_Installation_Guide.md"),
            ("현장 운영", "Protocol/Shelter_Ops_Deploy.xlsx"),
            ("CMS 배포 (참조)", "Shelter_CMS_Deploy/"),
        ],
        col_widths=[5, 10],
    )

    doc.add_heading("부록 C — 개정", level=1)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ("본 문서", f"rev {REV}"),
            ("규격 개정", "Shelter_Protocol.md 버전 이력"),
            ("엑셀 재생성", "python _build_shelter_protocol_xlsx.py"),
            ("Word 재생성", "python _build_개발가이드_docx.py"),
        ],
        col_widths=[4, 12],
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f"Smart Shelter 개발 가이드 rev {REV} — MOS_version")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    fr.font.name = FONT_KO
    fr._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)


def main():
    doc = Document()
    set_doc_defaults(doc)
    build_cover(doc)
    build_body(doc)
    doc.save(OUT)
    doc.save(OUT_V1)
    print(f"Saved: {OUT}")
    print(f"Saved: {OUT_V1}")
    if MD_SRC.exists():
        print(f"Source MD: {MD_SRC}")


if __name__ == "__main__":
    main()
